import json
import sys

import docx
import tqdm


def make_spans(color_map, characters):
    result = []
    last_span_lst = None
    current_text = []
    for span_lst, c in zip(color_map, characters):
        if last_span_lst is None or last_span_lst == span_lst:
            current_text.append(c)
            last_span_lst = span_lst
        else:
            assert last_span_lst is not None and len(current_text) > 0
            result.append((last_span_lst, "".join(current_text)))
            last_span_lst = span_lst
            current_text = [c]
    else:
        # I really should have something here
        assert last_span_lst is not None and len(current_text) > 0
        result.append((last_span_lst, "".join(current_text)))
    return result


def index_fonts():
    palette = []
    with open("palette.txt") as f:
        for line in f:
            line = line.strip()
            if not line.startswith("#"):
                continue
            palette.append(line)
    # print(palette)
    assert len(palette) == len(set(palette))

    rgb_colors = []
    for colorspec in palette:
        color = docx.shared.RGBColor.from_string(colorspec[1:].upper())
        rgb_colors.append(color)

    return rgb_colors


rgb_colors = index_fonts()


def para2txt(p, p_idx, doc):
    total_len = 0

    pgraph = doc.add_paragraph()
    t = f"Text number {p_idx}"
    total_len += len(t)+1
    pgraph.add_run(t).bold = True

    pgraph = doc.add_paragraph('')

    ctx = p["context"]
    color_map = [[] for _ in range(len(ctx))]

    for qa in p["qas"]:
        q = qa["question"]
        q_id = qa["id"]
        for a_idx, a in enumerate(qa["answers"]+qa.get("plausible_answers", [])):
            atext = a["text"]
            a_char_idx = int(a["answer_start"])
            assert atext == ctx[a_char_idx:a_char_idx+len(atext)]
            for lst in color_map[a_char_idx:a_char_idx+len(atext)]:
                lst.append(q_id+"_"+str(a_idx))
            # print("ATEXT",atext)
            # print("CTXT ",ctx[aidx:aidx+len(atext)])
            # print()

    # make a lookup table such that each unique overlap of answer ids has a color of its own whew
    color_lookup = {"": -1}
    for answer_list in color_map:  # this is a list of all questions which overlap as "questionid_answeridx" strings
        answer_list = "+".join(answer_list)  # make it a single string
        color_lookup.setdefault(answer_list, len(color_lookup))

    # print(color_lookup)
    # colors=index_fonts()

    spans = make_spans(color_map, ctx)
    for span_list, txt in spans:
        r = pgraph.add_run(txt)
        font_idx = color_lookup.get("+".join(span_list), None)
        if font_idx >= 0:
            r.font.color.rgb = rgb_colors[font_idx]

    total_len += len(ctx)+1

    question_list = []
    for q_idx, qa in enumerate(p["qas"]):
        q = qa["question"]
        question_list.append(qa["id"])
        pgraph = doc.add_paragraph("")
        t = f"Question {q_idx}"
        total_len += len(t)
        pgraph.add_run(t).bold = True
        pgraph = doc.add_paragraph(q)
        total_len += len(q)+2

    return total_len, color_lookup, question_list


if __name__ == "__main__":

    doc = docx.Document()

    total_len = 0
    file_counter = 0

    sizes = []
    d_idx = 0
    with open("squad2-en/meta.jsonl", "wt") as meta:
        for fname in sys.argv[1:]:
            with open(fname, "rt") as f:
                data = json.load(f)["data"]
                for d in tqdm.tqdm(data):
                    docmeta = {
                        "file": fname, "title": d["title"], "sequence_idx": d_idx, "paragraphs": []}
                    title = d["title"]
                    pgraph = doc.add_paragraph("")
                    r = pgraph.add_run(f"Document number {d_idx}")
                    r.bold = True
                    r.underline = True
                    ps = d["paragraphs"]
                    for p_idx, p in enumerate(ps):
                        l, color_map, question_list = para2txt(p, p_idx, doc)
                        docmeta["paragraphs"].append(
                            (p_idx, color_map, question_list))
                        total_len += l
                    print(json.dumps(docmeta, sort_keys=True,
                          ensure_ascii=False), file=meta, flush=True)
                    doc.add_page_break()
                    d_idx += 1
                    if total_len > 900000:  # document full!
                        doc.save(f"squad2-en/squad2_{file_counter:03d}.docx")
                        file_counter += 1
                        total_len = 0
                        doc = docx.Document()
        else:
            doc.save(f"squad2-en/squad2_{file_counter:03d}.docx")
